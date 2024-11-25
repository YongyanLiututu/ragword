import os
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_mistralai.embeddings import MistralAIEmbeddings
from langchain_mistralai.chat_models import ChatMistralAI
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from docx import Document  # 用于加载 Word 文档
import faiss  # FAISS 库

# 加载环境变量
load_dotenv()
mistral_api_key = "S5lLgc9Qkmx1eOTspsBXahjxBEsK9EZ6"


class RAG_Word:
    """使用 LangChain 和 MistralAI 模型处理 Word 文档的问答系统类。"""

    def __init__(self, chat_model="open-mixtral-8x7b", emb_model="mistral-embed"):
        """初始化问答系统，包括嵌入模型和聊天模型。"""
        self.model = ChatMistralAI(model=chat_model, temperature=0, mistral_api_key=mistral_api_key)
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
        self.embeddings = MistralAIEmbeddings(model=emb_model, mistral_api_key=mistral_api_key)
        self.system_prompt = (
            """
            <s> [INST] 你是一名问答任务助手。请使用以下上下文回答问题。如果不知道答案，请直接说你不知道。
            答案最多不超过五句话，简明扼要。 [/INST] </s>
            [INST] 上下文: {context} [/INST]
            """
        )
        self.vector_store = None
        self.retriever = None
        self.rag_chain = None

    def load_word_document(self, file_path):
        """加载 Word 文档中的文本内容。"""
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text

    def read_directory(self, directory_path: str):
        """处理指定目录下的所有 Word 文档，并加载或重建向量存储。"""
        print("正在构建向量存储...")
        word_files = [f for f in os.listdir(directory_path) if f.endswith('.docx')]
        all_docs = []

        for word_file in word_files:
            full_path = os.path.join(directory_path, word_file)
            text = self.load_word_document(full_path)

            # 使用 Document 对象而不是字典
            all_docs.append(Document(page_content=text))

        chunks = self.text_splitter.split_documents(all_docs)

        # 使用 FAISS 初始化向量存储
        self.vector_store = FAISS.from_documents(chunks, embedding=self.embeddings)
        self.retriever = self.vector_store.as_retriever(search_kwargs={"k": 5})

        # 设置处理链用于处理查询
        self.prompt = ChatPromptTemplate.from_messages(
            [
                ("system", self.system_prompt),
                ("human", "{input}"),
            ]
        )
        question_answer_chain = create_stuff_documents_chain(self.model, self.prompt)
        self.rag_chain = create_retrieval_chain(self.retriever, question_answer_chain)

    def ask(self, query: str):
        """根据输入问题提供答案及相关上下文。"""
        if not self.rag_chain:
            return "向量数据库未构建。", []
        answer = self.rag_chain.invoke({"input": query})
        return answer

    def clear(self):
        """清除存储的数据并重置系统。"""
        self.vector_store = None
        self.retriever = None
        self.rag_chain = None
