import os
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_mistralai.embeddings import MistralAIEmbeddings
from langchain_mistralai.chat_models import ChatMistralAI
from langchain_core.prompts import ChatPromptTemplate
from langchain.schema import Document  # 用于构造 LangChain 的文档对象
from docx import Document as WordDocument  # 用于加载 Word 文档
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import warnings

# 忽略特定警告（可选）
warnings.filterwarnings("ignore", message="Could not download mistral tokenizer")

# 加载环境变量
load_dotenv()
mistral_api_key = os.getenv("MISTRAL_API_KEY", "api")


# 文档预处理 Agent
class DocumentPreprocessingAgent:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)

    def process_document(self, file_path):
        """加载并分块 Word 文档。"""
        try:
            doc = WordDocument(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if not text:
                raise ValueError(f"文档 {file_path} 是空的。")
            chunks = self.text_splitter.split_text(text)
            return [Document(page_content=chunk, metadata={"source": file_path}) for chunk in chunks]
        except Exception as e:
            raise ValueError(f"处理文档失败: {e}")

    def process_directory(self, directory_path):
        """处理目录中的所有 Word 文档。"""
        word_files = [f for f in os.listdir(directory_path) if f.endswith('.docx')]
        all_chunks = []
        for word_file in word_files:
            file_path = os.path.join(directory_path, word_file)
            chunks = self.process_document(file_path)
            all_chunks.extend(chunks)
        return all_chunks


# 向量存储 Agent
class VectorStoreAgent:
    def __init__(self, embeddings):
        self.vector_store = None
        self.embeddings = embeddings

    def build_store(self, documents):
        self.vector_store = FAISS.from_documents(documents, embedding=self.embeddings)

    def retrieve(self, query, k=5):
        if not self.vector_store:
            raise ValueError("向量存储未初始化。")
        retriever = self.vector_store.as_retriever(search_kwargs={"k": k})
        return retriever.get_relevant_documents(query)


# 问答 Agent
class QAAgent:
    def __init__(self, chat_model, system_prompt):
        self.model = chat_model
        self.system_prompt = system_prompt

    def generate_answer(self, retriever, query):
        # 从 retriever 获取相关文档
        documents = retriever(query)
        context = "\n".join([doc.page_content for doc in documents])
        
        # 构造符合格式的消息
        messages = [
            {"role": "system", "content": self.system_prompt},
            {"role": "user", "content": f"上下文: {context}\n问题: {query}"}
        ]

        # 检查消息格式是否符合预期
        if not all(isinstance(msg, dict) and "role" in msg and "content" in msg for msg in messages):
            raise ValueError(f"消息格式不正确: {messages}")

        # 调试输出消息结构
        print("构造的消息结构: ", messages)

        # 调用模型生成答案
        return self.model.generate(messages)






# 推理链 Agent
class ReasoningAgent:
    def __init__(self, model):
        self.model = model

    def reason(self, context, question, depth=3):
        """通过多轮推理生成答案。"""
        steps = []
        current_question = question
        for _ in range(depth):
            response = self.model.generate(f"上下文：{context}\n\n问题：{current_question}\n回答：")
            steps.append(response)
            current_question = f"基于之前的回答：{response}，还有什么问题需要解决？"
        return "\n".join(steps)


# 生成 Agent
class GenerationAgent:
    def __init__(self, chat_model):
        self.model = chat_model

    def generate_summary(self, context):
        """根据上下文生成摘要。"""
        return self.model.generate(f"根据以下内容生成摘要：\n{context}")


# 任务协调 Agent
class TaskCoordinatorAgent:
    def __init__(self):
        self.agents = {}

    def register_agent(self, task_name, agent):
        """注册子 Agent。"""
        self.agents[task_name] = agent

    def execute(self, task_name, *args, **kwargs):
        """执行任务。"""
        if task_name not in self.agents:
            raise ValueError(f"任务未注册: {task_name}")
        return self.agents[task_name](*args, **kwargs)


# 主系统类
class RAGSystem:
    def __init__(self):
        # 初始化各个 Agent
        self.preprocessor = DocumentPreprocessingAgent()
        self.embeddings = MistralAIEmbeddings(model="mistral-embed", mistral_api_key=mistral_api_key)
        self.vector_store = VectorStoreAgent(self.embeddings)
        self.chat_model = ChatMistralAI(model="mistral-small-latest", temperature=0, mistral_api_key=mistral_api_key)
        self.qa_agent = QAAgent(
            self.chat_model,
            system_prompt="<s>[INST] 你是一名问答助手。使用上下文内容回答问题... [/INST]</s>",
        )
        self.reasoning_agent = ReasoningAgent(self.chat_model)
        self.generation_agent = GenerationAgent(self.chat_model)

        # 初始化任务协调 Agent
        self.coordinator = TaskCoordinatorAgent()
        self.coordinator.register_agent("preprocess", self.preprocessor.process_directory)
        self.coordinator.register_agent("build_store", self.vector_store.build_store)
        self.coordinator.register_agent("retrieve", self.vector_store.retrieve)
        self.coordinator.register_agent("qa", self.qa_agent.generate_answer)
        self.coordinator.register_agent("reasoning", self.reasoning_agent.reason)
        self.coordinator.register_agent("generate_summary", self.generation_agent.generate_summary)

    def process_documents(self, directory_path):
        """处理目录中的文档并构建向量存储。"""
        chunks = self.coordinator.execute("preprocess", directory_path)
        self.coordinator.execute("build_store", chunks)

    def ask_question(self, query):
        """问答流程。"""
        retriever = lambda q: self.coordinator.execute("retrieve", q)
        return self.coordinator.execute("qa", retriever, query)

    def perform_reasoning(self, context, question, depth=3):
        """推理流程。"""
        return self.coordinator.execute("reasoning", context, question, depth)

    def generate_summary(self, context):
        """生成摘要。"""
        return self.coordinator.execute("generate_summary", context)
